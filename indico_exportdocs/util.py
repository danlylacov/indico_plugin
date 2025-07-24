from docx import Document
from indico.modules.events.models.events import Event
from indico.modules.events.papers.models.revisions import PaperRevisionState
from io import BytesIO

def generate_docx_list(event_id):
    event = Event.get(event_id)
    doc = Document()
    doc.add_heading(f'Список докладов для "{event.title}"', 0)
    for session in event.sessions:
        doc.add_heading(f'Заседание: {session.title}', level=1)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Название доклада'
        hdr_cells[1].text = 'Докладчики'
        hdr_cells[2].text = 'Время'
        for c in session.contributions:
            row = table.add_row().cells
            row[0].text = c.title or ''
            row[1].text = ', '.join([p.full_name for p in c.person_links if p.is_speaker])
            row[2].text = c.start_dt.strftime('%Y-%m-%d %H:%M') if c.start_dt else ''
    # Доклады вне сессий
    other_contribs = [c for c in event.contributions if not c.session]
    if other_contribs:
        doc.add_heading('Доклады вне сессий', level=1)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Название доклада'
        hdr_cells[1].text = 'Докладчики'
        hdr_cells[2].text = 'Время'
        for c in other_contribs:
            row = table.add_row().cells
            row[0].text = c.title or ''
            row[1].text = ', '.join([p.full_name for p in c.person_links if p.is_speaker])
            row[2].text = c.start_dt.strftime('%Y-%m-%d %H:%M') if c.start_dt else ''
    f = BytesIO()
    doc.save(f)
    return f.getvalue()

def generate_docx_report(event_id):
    event = Event.get(event_id)
    doc = Document()
    doc.add_heading(f'Отчет о проведении конференции "{event.title}"', 0)
    for session in event.sessions:
        doc.add_heading(f'Заседание: {session.title}', level=1)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Название доклада'
        hdr_cells[1].text = 'Докладчики'
        hdr_cells[2].text = 'Время'
        hdr_cells[3].text = 'Отметка о выступлении'
        for c in session.contributions:
            row = table.add_row().cells
            row[0].text = c.title or ''
            row[1].text = ', '.join([p.full_name for p in c.person_links if p.is_speaker])
            row[2].text = c.start_dt.strftime('%Y-%m-%d %H:%M') if c.start_dt else ''
            row[3].text = ''  # Пустой столбец для ручного заполнения
    # Доклады вне сессий
    other_contribs = [c for c in event.contributions if not c.session]
    if other_contribs:
        doc.add_heading('Доклады вне сессий', level=1)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Название доклада'
        hdr_cells[1].text = 'Докладчики'
        hdr_cells[2].text = 'Время'
        hdr_cells[3].text = 'Отметка о выступлении'
        for c in other_contribs:
            row = table.add_row().cells
            row[0].text = c.title or ''
            row[1].text = ', '.join([p.full_name for p in c.person_links if p.is_speaker])
            row[2].text = c.start_dt.strftime('%Y-%m-%d %H:%M') if c.start_dt else ''
            row[3].text = ''
    f = BytesIO()
    doc.save(f)
    return f.getvalue()

def generate_docx_papers(event_id):
    event = Event.get(event_id)
    doc = Document()
    doc.add_heading(f'Список предоставленных к публикации статей для "{event.title}"', 0)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Название статьи'
    hdr_cells[1].text = 'Авторы'
    hdr_cells[2].text = 'ID статьи'
    for c in event.contributions:
        rev = c._accepted_paper_revision
        if rev and rev.state == PaperRevisionState.accepted:
            row = table.add_row().cells
            row[0].text = c.title or ''
            row[1].text = ', '.join([p.full_name for p in c.person_links if p.is_speaker])
            row[2].text = str(rev.id)
    f = BytesIO()
    doc.save(f)
    return f.getvalue()
